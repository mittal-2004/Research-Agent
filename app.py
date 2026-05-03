import streamlit as st
import time
import io
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from agents import build_reader_agent, build_search_agent, writer_chain, critic_chain
 
# ── Page config ───────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="Synthetix · Deep Research Engine",
    page_icon="◈",
    layout="wide",
    initial_sidebar_state="collapsed",
)
 
# ── CSS ───────────────────────────────────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Space+Grotesk:wght@300;400;500;600;700&family=Syne:wght@400;500;600;700;800&family=JetBrains+Mono:wght@300;400;500&display=swap');
 
:root {
    --bg:           #f0f4ff;
    --bg2:          #ffffff;
    --bg3:          #f7f9ff;
    --bg4:          #eef1fc;
    --border:       #d4dbf5;
    --border2:      #c2ccee;
    --text:         #0f1630;
    --text-dim:     #4a5380;
    --text-muted:   #8a97c4;
    --indigo:       #4f46e5;
    --indigo-light: rgba(79,70,229,0.1);
    --indigo-glow:  rgba(79,70,229,0.2);
    --pink:         #ec4899;
    --pink-light:   rgba(236,72,153,0.1);
    --cyan:         #06b6d4;
    --amber:        #f59e0b;
    --emerald:      #10b981;
    --rose:         #f43f5e;
    --purple:       #8b5cf6;
}
 
*, *::before, *::after { box-sizing: border-box; margin: 0; }
 
html, body, [class*="css"] {
    font-family: 'Space Grotesk', sans-serif;
    color: var(--text);
}
 
.stApp {
    background: var(--bg);
    background-image:
        radial-gradient(ellipse 60% 50% at 0% 0%, rgba(79,70,229,0.12) 0%, transparent 55%),
        radial-gradient(ellipse 50% 40% at 100% 100%, rgba(236,72,153,0.1) 0%, transparent 50%),
        radial-gradient(ellipse 40% 30% at 50% 50%, rgba(6,182,212,0.05) 0%, transparent 60%);
}
 
#MainMenu, footer, header { visibility: hidden; }
.block-container { padding: 2rem 3rem 5rem; max-width: 1300px; }
 
/* ── TOPBAR ── */
.topbar {
    display: flex;
    align-items: center;
    justify-content: space-between;
    padding: 1.1rem 0 1.5rem;
    border-bottom: 2px solid var(--border);
    margin-bottom: 3rem;
}
.brand { display: flex; align-items: center; gap: 0.7rem; }
.brand-icon {
    width: 38px; height: 38px;
    background: linear-gradient(135deg, var(--indigo), var(--pink));
    border-radius: 10px;
    display: flex; align-items: center; justify-content: center;
    font-size: 1.1rem; color: #fff;
    box-shadow: 0 4px 14px rgba(79,70,229,0.35);
}
.brand-name {
    font-family: 'Syne', sans-serif;
    font-size: 1.5rem;
    font-weight: 800;
    color: var(--text);
    letter-spacing: -0.04em;
}
.brand-name span { color: var(--indigo); }
.topbar-tag {
    font-family: 'JetBrains Mono', monospace;
    font-size: 0.62rem;
    letter-spacing: 0.15em;
    text-transform: uppercase;
    color: var(--indigo);
    background: var(--indigo-light);
    border: 1px solid rgba(79,70,229,0.25);
    padding: 0.3rem 0.75rem;
    border-radius: 20px;
    font-weight: 500;
}
 
/* ── HERO ── */
.hero { text-align: center; padding: 0.5rem 0 2.5rem; }
.hero-kicker {
    font-family: 'JetBrains Mono', monospace;
    font-size: 0.68rem;
    letter-spacing: 0.22em;
    text-transform: uppercase;
    color: var(--pink);
    margin-bottom: 1.1rem;
    display: flex; align-items: center; justify-content: center; gap: 0.6rem;
    font-weight: 500;
}
.hero-kicker-dot {
    width: 5px; height: 5px; border-radius: 50%;
    background: var(--pink); display: inline-block;
}
.hero h1 {
    font-family: 'Syne', sans-serif;
    font-size: clamp(2.6rem, 6vw, 4.5rem);
    font-weight: 800;
    line-height: 1.08;
    letter-spacing: -0.04em;
    color: var(--text);
    margin-bottom: 1rem;
}
.hero h1 .grad {
    background: linear-gradient(135deg, var(--indigo), var(--pink), var(--cyan));
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
    background-clip: text;
}
.hero-desc {
    font-size: 1.05rem;
    font-weight: 400;
    color: var(--text-dim);
    max-width: 500px;
    margin: 0 auto;
    line-height: 1.75;
}
 
/* ── DIVIDER ── */
.glow-divider {
    height: 2px;
    background: linear-gradient(90deg, transparent, var(--indigo), var(--pink), var(--cyan), transparent);
    margin: 1.5rem 0 2.5rem;
    border-radius: 2px;
    opacity: 0.4;
}
 
/* ── INPUT SHELL ── */
.input-shell {
    background: var(--bg2);
    border: 1.5px solid var(--border2);
    border-radius: 18px;
    padding: 1.8rem 2rem;
    position: relative;
    overflow: hidden;
    margin-bottom: 0.5rem;
    box-shadow: 0 4px 24px rgba(79,70,229,0.06), 0 1px 4px rgba(0,0,0,0.04);
}
.input-shell::before {
    content: '';
    position: absolute; top: 0; left: 0; right: 0; height: 3px;
    background: linear-gradient(90deg, var(--indigo), var(--pink), var(--cyan));
    border-radius: 18px 18px 0 0;
}
.input-shell-label {
    font-family: 'JetBrains Mono', monospace;
    font-size: 0.65rem; letter-spacing: 0.18em;
    text-transform: uppercase; color: var(--indigo);
    margin-bottom: 0.8rem; font-weight: 500;
    display: flex; align-items: center; gap: 0.5rem;
}
.input-shell-label::before { content: '▸'; font-size: 0.7rem; }
 
/* Streamlit input overrides */
.stTextInput > label { display: none !important; }
.stTextInput > div > div > input {
    background: var(--bg3) !important;
    border: 1.5px solid var(--border2) !important;
    border-radius: 10px !important;
    color: var(--text) !important;
    font-family: 'Space Grotesk', sans-serif !important;
    font-size: 1rem !important;
    font-weight: 400 !important;
    padding: 0.85rem 1.1rem !important;
    caret-color: var(--indigo) !important;
    transition: border-color 0.2s, box-shadow 0.2s !important;
}
.stTextInput > div > div > input::placeholder { color: var(--text-muted) !important; }
.stTextInput > div > div > input:focus {
    border-color: var(--indigo) !important;
    box-shadow: 0 0 0 4px var(--indigo-glow) !important;
    outline: none !important;
}
 
/* Button */
.stButton > button {
    background: linear-gradient(135deg, var(--indigo) 0%, var(--purple) 50%, var(--pink) 100%) !important;
    color: #fff !important;
    font-family: 'Syne', sans-serif !important;
    font-weight: 700 !important;
    font-size: 0.95rem !important;
    letter-spacing: 0.02em !important;
    border: none !important;
    border-radius: 10px !important;
    padding: 0.8rem 1.8rem !important;
    width: 100% !important;
    cursor: pointer !important;
    box-shadow: 0 4px 20px rgba(79,70,229,0.35) !important;
    transition: transform 0.15s, box-shadow 0.2s !important;
}
.stButton > button:hover {
    transform: translateY(-2px) !important;
    box-shadow: 0 8px 32px rgba(79,70,229,0.45) !important;
}
.stButton > button:active { transform: translateY(0) !important; }
 
/* Chips */
.chips-row {
    display: flex; flex-wrap: wrap;
    align-items: center; gap: 0.45rem;
    margin-top: 1.1rem;
}
.chips-label {
    font-family: 'JetBrains Mono', monospace;
    font-size: 0.6rem; letter-spacing: 0.12em;
    color: var(--text-muted); text-transform: uppercase;
}
.chip {
    background: var(--bg4);
    border: 1.5px solid var(--border);
    border-radius: 20px;
    padding: 0.22rem 0.75rem;
    font-size: 0.73rem;
    color: var(--text-dim);
    font-family: 'Space Grotesk', sans-serif;
    font-weight: 500;
    cursor: default;
    transition: border-color 0.2s;
}
.chip:hover { border-color: var(--indigo); color: var(--indigo); }
 
/* ── PIPELINE ── */
.pipeline-heading {
    font-family: 'JetBrains Mono', monospace;
    font-size: 0.65rem; letter-spacing: 0.2em;
    text-transform: uppercase; color: var(--text-muted);
    margin-bottom: 1.1rem; padding-left: 0.2rem; font-weight: 500;
}
 
.step-card {
    background: var(--bg2);
    border: 1.5px solid var(--border);
    border-radius: 12px;
    padding: 1rem 1.2rem;
    margin-bottom: 0.65rem;
    position: relative; overflow: hidden;
    transition: border-color 0.3s, background 0.3s, box-shadow 0.3s;
    box-shadow: 0 1px 4px rgba(0,0,0,0.04);
}
.step-card::before {
    content: '';
    position: absolute; left: 0; top: 0; bottom: 0; width: 3px;
    background: var(--border2); border-radius: 3px 0 0 3px;
    transition: background 0.3s;
}
.step-card.active {
    border-color: rgba(79,70,229,0.4);
    background: linear-gradient(135deg, rgba(79,70,229,0.04), rgba(139,92,246,0.03));
    box-shadow: 0 4px 16px rgba(79,70,229,0.1);
}
.step-card.active::before { background: linear-gradient(180deg, var(--indigo), var(--purple)); }
.step-card.done  { border-color: rgba(16,185,129,0.35); background: rgba(16,185,129,0.04); }
.step-card.done::before  { background: var(--emerald); }
 
.step-inner { display: flex; align-items: center; gap: 0.8rem; }
.step-num {
    font-family: 'JetBrains Mono', monospace;
    font-size: 0.62rem; font-weight: 500;
    color: var(--text-muted); letter-spacing: 0.08em; min-width: 24px;
}
.step-card.active .step-num { color: var(--indigo); }
.step-card.done .step-num   { color: var(--emerald); }
 
.step-title {
    font-family: 'Space Grotesk', sans-serif;
    font-size: 0.9rem; font-weight: 600;
    color: var(--text-dim); flex: 1;
}
.step-card.active .step-title { color: var(--text); }
.step-card.done .step-title   { color: var(--text-dim); }
 
.step-desc {
    font-size: 0.74rem; color: var(--text-muted);
    margin-top: 0.3rem; padding-left: 2.3rem; font-weight: 400;
}
 
.step-status {
    font-family: 'JetBrains Mono', monospace;
    font-size: 0.6rem; letter-spacing: 0.1em;
    padding: 0.2rem 0.55rem; border-radius: 20px; font-weight: 500;
}
.s-wait { color: var(--text-muted); background: var(--bg4); }
.s-run  { color: var(--indigo); background: var(--indigo-light); animation: flicker 1.4s infinite; border: 1px solid rgba(79,70,229,0.2); }
.s-done { color: var(--emerald); background: rgba(16,185,129,0.1); border: 1px solid rgba(16,185,129,0.2); }
 
@keyframes flicker { 0%,100%{opacity:1} 50%{opacity:0.4} }
 
/* ── STATUS BAR ── */
.status-bar {
    display: flex; align-items: center; gap: 0.7rem;
    background: var(--indigo-light);
    border: 1.5px solid rgba(79,70,229,0.25);
    border-left: 4px solid var(--indigo);
    border-radius: 0 10px 10px 0;
    padding: 0.7rem 1.2rem;
    font-family: 'JetBrains Mono', monospace;
    font-size: 0.72rem; color: var(--indigo); font-weight: 500;
    margin: 1rem 0;
}
.status-bar.done {
    border-left-color: var(--emerald);
    background: rgba(16,185,129,0.08);
    border-color: rgba(16,185,129,0.25);
    color: var(--emerald);
}
.pulse {
    width: 8px; height: 8px; border-radius: 50%;
    background: var(--indigo); flex-shrink: 0;
    animation: pulse 1.1s infinite;
}
@keyframes pulse { 0%,100%{opacity:1;transform:scale(1)} 50%{opacity:0.3;transform:scale(0.5)} }
 
/* ── RESULTS ── */
.results-heading {
    font-family: 'Syne', sans-serif;
    font-size: 1.4rem; font-weight: 700;
    letter-spacing: -0.03em; color: var(--text);
    margin: 2.5rem 0 1.25rem;
}
 
.raw-content {
    font-size: 0.83rem; line-height: 1.8;
    color: var(--text-dim); white-space: pre-wrap;
    font-family: 'Space Grotesk', sans-serif; font-weight: 400;
    padding: 1.2rem;
    max-height: 280px; overflow-y: auto;
    scrollbar-width: thin;
    scrollbar-color: var(--border2) transparent;
}
 
details summary {
    font-family: 'JetBrains Mono', monospace !important;
    font-size: 0.7rem !important;
    letter-spacing: 0.12em !important;
    color: var(--text-dim) !important;
    text-transform: uppercase !important;
    cursor: pointer !important;
    font-weight: 500 !important;
}
 
/* Report shell */
.report-shell {
    background: var(--bg2);
    border: 1.5px solid var(--border2);
    border-radius: 16px;
    padding: 2rem 2.4rem;
    margin-bottom: 0.5rem;
    position: relative; overflow: hidden;
    box-shadow: 0 4px 24px rgba(245,158,11,0.08);
}
.report-shell::before {
    content: '';
    position: absolute; top: 0; left: 0; right: 0; height: 3px;
    background: linear-gradient(90deg, var(--amber), var(--pink), var(--purple));
}
.panel-label {
    font-family: 'JetBrains Mono', monospace;
    font-size: 0.65rem; letter-spacing: 0.18em;
    text-transform: uppercase;
    margin-bottom: 1.2rem; padding-bottom: 0.8rem;
    display: flex; align-items: center; gap: 0.5rem;
    font-weight: 500;
}
.panel-label.amber {
    color: var(--amber);
    border-bottom: 1.5px solid rgba(245,158,11,0.2);
}
.panel-label.rose {
    color: var(--rose);
    border-bottom: 1.5px solid rgba(244,63,94,0.15);
}
 
/* Critic shell */
.critic-shell {
    background: var(--bg2);
    border: 1.5px solid var(--border2);
    border-radius: 16px;
    padding: 2rem 2.4rem;
    position: relative; overflow: hidden;
    box-shadow: 0 4px 24px rgba(244,63,94,0.06);
}
.critic-shell::before {
    content: '';
    position: absolute; top: 0; left: 0; right: 0; height: 3px;
    background: linear-gradient(90deg, var(--rose), var(--pink), var(--purple));
}
 
/* Download btn */
.stDownloadButton > button {
    background: linear-gradient(135deg, rgba(79,70,229,0.08), rgba(139,92,246,0.06)) !important;
    border: 1.5px solid rgba(79,70,229,0.3) !important;
    color: var(--indigo) !important;
    font-family: 'JetBrains Mono', monospace !important;
    font-size: 0.72rem !important;
    letter-spacing: 0.1em !important;
    border-radius: 10px !important;
    padding: 0.55rem 1.3rem !important;
    margin-top: 1rem !important;
    font-weight: 500 !important;
    transition: all 0.2s !important;
}
.stDownloadButton > button:hover {
    background: linear-gradient(135deg, rgba(79,70,229,0.15), rgba(139,92,246,0.1)) !important;
    border-color: var(--indigo) !important;
    box-shadow: 0 4px 14px rgba(79,70,229,0.2) !important;
}
 
/* Misc */
.stSpinner > div { border-top-color: var(--indigo) !important; }
 
.site-footer {
    margin-top: 4rem; padding-top: 1.5rem;
    border-top: 1.5px solid var(--border);
    display: flex; align-items: center; justify-content: space-between;
}
.footer-txt {
    font-family: 'JetBrains Mono', monospace;
    font-size: 0.6rem; color: var(--text-muted);
    letter-spacing: 0.1em; text-transform: uppercase; font-weight: 400;
}
 
[data-testid="column"] { padding: 0 0.5rem !important; }
[data-testid="column"]:first-child { padding-left: 0 !important; }
[data-testid="column"]:last-child  { padding-right: 0 !important; }
</style>
""", unsafe_allow_html=True)
 
 
# ── DOCX generation helper ────────────────────────────────────────────────────
def generate_docx(topic: str, report: str, critic: str) -> bytes:
    doc = DocxDocument()
 
    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
 
    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("SYNTHETIX RESEARCH REPORT")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)  # indigo
 
    # Topic subtitle
    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_para.add_run(topic)
    sub_run.font.size = Pt(13)
    sub_run.font.color.rgb = RGBColor(0x8B, 0x5C, 0xF6)  # purple
    sub_run.italic = True
 
    # Timestamp
    ts_para = doc.add_paragraph()
    ts_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ts_run = ts_para.add_run(f"Generated: {time.strftime('%Y-%m-%d %H:%M UTC')}")
    ts_run.font.size = Pt(9)
    ts_run.font.color.rgb = RGBColor(0x8A, 0x97, 0xC4)
 
    doc.add_paragraph()  # spacer
 
    # Divider heading
    h = doc.add_heading("Research Report", level=1)
    h.runs[0].font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)
 
    # Report content — split on markdown headings
    for line in report.splitlines():
        stripped = line.strip()
        if stripped.startswith("### "):
            p = doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            p = doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            p = doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(stripped[2:], style="List Bullet")
        elif stripped == "":
            doc.add_paragraph()
        else:
            doc.add_paragraph(stripped)
 
    doc.add_page_break()
 
    # Critic section
    h2 = doc.add_heading("Critic Review", level=1)
    h2.runs[0].font.color.rgb = RGBColor(0xF4, 0x3F, 0x5E)
 
    for line in critic.splitlines():
        stripped = line.strip()
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif stripped == "":
            doc.add_paragraph()
        else:
            doc.add_paragraph(stripped)
 
    # Footer note
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer_para.add_run("◈ Powered by Synthetix · Deep Research Engine")
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(0x8A, 0x97, 0xC4)
    fr.italic = True
 
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
 
 
# ── Step card helper ──────────────────────────────────────────────────────────
def step_card(num, title, desc, state):
    status_map = {
        "waiting": ("IDLE",    "s-wait", ""),
        "running": ("RUNNING", "s-run",  "active"),
        "done":    ("✓ DONE",  "s-done", "done"),
    }
    label, scls, ccls = status_map.get(state, ("", "", ""))
    st.markdown(f"""
    <div class="step-card {ccls}">
      <div class="step-inner">
        <span class="step-num">{num}</span>
        <span class="step-title">{title}</span>
        <span class="step-status {scls}">{label}</span>
      </div>
      <div class="step-desc">{desc}</div>
    </div>
    """, unsafe_allow_html=True)
 
 
# ── Session state init ────────────────────────────────────────────────────────
for k, v in [("results", {}), ("running", False), ("done", False)]:
    if k not in st.session_state:
        st.session_state[k] = v
 
 
# ── Topbar ────────────────────────────────────────────────────────────────────
st.markdown("""
<div class="topbar">
  <div class="brand">
    <div class="brand-icon">◈</div>
    <span class="brand-name">Synthe<span>tix</span></span>
  </div>
  <span class="topbar-tag">Deep Research Engine</span>
</div>
""", unsafe_allow_html=True)
 
 
# ── Hero ──────────────────────────────────────────────────────────────────────
st.markdown("""
<div class="hero">
  <div class="hero-kicker">
    <span class="hero-kicker-dot"></span>
    Multi-Agent AI System
    <span class="hero-kicker-dot"></span>
  </div>
  <h1>Research at<br><span class="grad">machine speed.</span></h1>
  <p class="hero-desc">
    Four specialized agents collaborate — searching, scraping, writing, and critiquing —
    to deliver a polished research report in seconds.
  </p>
</div>
<div class="glow-divider"></div>
""", unsafe_allow_html=True)
 
 
# ── Two-column layout ─────────────────────────────────────────────────────────
col_left, col_gap, col_right = st.columns([5, 0.4, 4])
 
with col_left:
    st.markdown('<div class="input-shell"><div class="input-shell-label">Research Topic</div>', unsafe_allow_html=True)
    topic = st.text_input(
        "topic",
        placeholder="e.g. CRISPR gene editing breakthroughs 2025…",
        key="topic_input",
        label_visibility="collapsed",
    )
    run_btn = st.button("◈  Run Synthetix Pipeline", use_container_width=True)
    st.markdown("""
    <div class="chips-row">
      <span class="chips-label">Try →</span>
      <span class="chip">LLM agents 2025</span>
      <span class="chip">Fusion energy</span>
      <span class="chip">Mars colonization</span>
      <span class="chip">Quantum computing</span>
    </div>
    """, unsafe_allow_html=True)
    st.markdown('</div>', unsafe_allow_html=True)
 
with col_right:
    st.markdown('<div class="pipeline-heading">◈ Pipeline Status</div>', unsafe_allow_html=True)
 
    r = st.session_state.results
 
    def get_state(step):
        steps = ["search", "reader", "writer", "critic"]
        if step in r:
            return "done"
        if st.session_state.running:
            for k in steps:
                if k not in r:
                    return "running" if k == step else "waiting"
        return "waiting"
 
    step_card("01", "Search Agent",  "Gathers recent web information",        get_state("search"))
    step_card("02", "Reader Agent",  "Scrapes & extracts primary content",    get_state("reader"))
    step_card("03", "Writer Chain",  "Drafts the full structured report",     get_state("writer"))
    step_card("04", "Critic Chain",  "Reviews, scores and flags weak areas",  get_state("critic"))
 
 
# ── Trigger run ───────────────────────────────────────────────────────────────
if run_btn:
    if not topic.strip():
        st.warning("Please enter a research topic.")
    else:
        st.session_state.results = {}
        st.session_state.running = True
        st.session_state.done = False
        st.rerun()
 
 
# ── Pipeline execution ────────────────────────────────────────────────────────
if st.session_state.running and not st.session_state.done:
    results = {}
    topic_val = st.session_state.topic_input
 
    st.markdown('<div class="status-bar"><span class="pulse"></span>Pipeline running — this may take a moment…</div>',
                unsafe_allow_html=True)
 
    with st.spinner(""):
        # Step 1 — Search
        sa = build_search_agent()
        sr = sa.invoke({"messages": [("user", f"Find recent, reliable and detailed information about: {topic_val}")]})
        results["search"] = sr["messages"][-1].content
        st.session_state.results = dict(results)
 
        # Step 2 — Reader
        ra = build_reader_agent()
        rr = ra.invoke({"messages": [("user",
            f"Based on the following search results about '{topic_val}', "
            f"pick the most relevant URL and scrape it for deeper content.\n\n"
            f"Search Results:\n{results['search'][:800]}")]})
        results["reader"] = rr["messages"][-1].content
        st.session_state.results = dict(results)
 
        # Step 3 — Writer
        combined = f"SEARCH RESULTS:\n{results['search']}\n\nDETAILED SCRAPED CONTENT:\n{results['reader']}"
        results["writer"] = writer_chain.invoke({"topic": topic_val, "research": combined})
        st.session_state.results = dict(results)
 
        # Step 4 — Critic
        results["critic"] = critic_chain.invoke({"report": results["writer"]})
        st.session_state.results = dict(results)
 
    st.session_state.running = False
    st.session_state.done = True
    st.rerun()
 
 
# ── Results display ───────────────────────────────────────────────────────────
r = st.session_state.results
 
if r:
    st.markdown('<div class="glow-divider"></div>', unsafe_allow_html=True)
    st.markdown('<div class="results-heading">Output</div>', unsafe_allow_html=True)
 
    if st.session_state.done:
        st.markdown(
            '<div class="status-bar done">✓ &nbsp;All agents complete — report ready</div>',
            unsafe_allow_html=True,
        )
 
    # Raw collapsible outputs
    if "search" in r:
        with st.expander("▸ Search Agent · Raw Output"):
            st.markdown(f'<div class="raw-content">{r["search"]}</div>', unsafe_allow_html=True)
 
    if "reader" in r:
        with st.expander("▸ Reader Agent · Scraped Content"):
            st.markdown(f'<div class="raw-content">{r["reader"]}</div>', unsafe_allow_html=True)
 
    # Final report
    if "writer" in r:
        st.markdown("""
        <div class="report-shell">
          <div class="panel-label amber">◈ Final Research Report</div>
        """, unsafe_allow_html=True)
        st.markdown(r["writer"])
        st.markdown("</div>", unsafe_allow_html=True)
 
        # ── Word document download ──
        topic_val = st.session_state.get("topic_input", "report")
        critic_text = r.get("critic", "")
        docx_bytes = generate_docx(topic_val, r["writer"], critic_text)
        st.download_button(
            label="⬇  Download Report (.docx)",
            data=docx_bytes,
            file_name=f"synthetix_report_{int(time.time())}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
 
    # Critic
    if "critic" in r:
        st.markdown("""
        <div class="critic-shell">
          <div class="panel-label rose">◈ Critic Review</div>
        """, unsafe_allow_html=True)
        st.markdown(r["critic"])
        st.markdown("</div>", unsafe_allow_html=True)
 
 
# ── Footer ────────────────────────────────────────────────────────────────────
st.markdown("""
<div class="site-footer">
  <span class="footer-txt">◈ Synthetix · Deep Research Engine</span>
  <span class="footer-txt">Search → Scrape → Write → Critique</span>
</div>
""", unsafe_allow_html=True)